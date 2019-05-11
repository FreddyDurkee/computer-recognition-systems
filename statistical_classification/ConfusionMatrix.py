import getopt
import os
import sys

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches
from pandas.compat import StringIO
from sklearn.metrics import confusion_matrix
from sklearn.metrics import precision_score, accuracy_score, recall_score, f1_score
from sklearn.utils.multiclass import unique_labels


def extractLabelsFromFile(path):
    labels = []
    file = open(path, "r")
    for line in file.readlines():
        labels.append(line)
        print line
    return labels


def extractClassificationCases(pathToClassificationData):
    trueLabels = []
    predictedLabels = []
    file = open(pathToClassificationData, "r")
    for line in file.readlines():
        splitedLine = line.split(";")
        trueLabelsCase = splitedLine[0].split(",")
        predictedLabelsCase = splitedLine[1].split(",")
        for trueLab, predictedLab in zip(trueLabelsCase, predictedLabelsCase):
            trueLabels.append(trueLab)
            predictedLabels.append(predictedLab)
        trueLabels = [x.strip() for x in trueLabels]
        predictedLabels = [x.strip() for x in predictedLabels]
    return trueLabels, predictedLabels


def plot_confusion_matrix(y_true, y_pred, classes,
                          normalize=False,
                          title=None,
                          cmap=plt.cm.Blues):
    cm = confusion_matrix(y_true, y_pred)

    if normalize:
        cm = cm.astype('float') / cm.sum(axis=1)[:, np.newaxis]

    fig, ax = plt.subplots()
    im = ax.imshow(cm, interpolation='nearest', cmap=cmap)
    ax.figure.colorbar(im, ax=ax)
    # We want to show all ticks...
    ax.set(xticks=np.arange(cm.shape[1]),
           yticks=np.arange(cm.shape[0]),
           # ... and label them with the respective list entries
           xticklabels=classes, yticklabels=classes,
           title=title,
           ylabel='True label',
           xlabel='Predicted label')

    # Rotate the tick labels and set their alignment.
    plt.setp(ax.get_xticklabels(), rotation=45, ha="right",
             rotation_mode="anchor")

    # create text annotations.
    fmt = '.2f' if normalize else 'd'
    thresh = cm.max() / 2.
    for i in range(cm.shape[0]):
        for j in range(cm.shape[1]):
            ax.text(j, i, format(cm[i, j], fmt),
                    ha="center", va="center",
                    color="white" if cm[i, j] > thresh else "black")
    fig.tight_layout()
    return fig


def classificationQuality(trueLabels, predictedLabels, average='weighted'):
    PRECISION = precision_score(trueLabels, predictedLabels, average=average)
    RECALL = recall_score(trueLabels, predictedLabels, average=average)
    F1 = f1_score(trueLabels, predictedLabels, average=average)
    ACC = accuracy_score(trueLabels, predictedLabels)
    return PRECISION, RECALL, F1, ACC


def plotClassificationQualityForSingleLabel(listOfTrueLabels, listOfPredictedLabels, listOfK):
    precissionList = []
    recallList = []
    f1List = []
    kList = []
    for trueLabels, predictedLabels, k in zip(listOfTrueLabels, listOfPredictedLabels, listOfK):
        precision, recall, f1, acc = classificationQuality(trueLabels, predictedLabels, average='weighted')
        precissionList.append(precision)
        recallList.append(recall)
        f1List.append(f1)
        kList.append(k)
    order = np.argsort(kList)
    precissionList = np.array(precissionList)[order]
    recallList = np.array(recallList)[order]
    f1List = np.array(f1List)[order]
    kList = np.array(kList)[order]

    plt.plot(kList, precissionList, label="Precision", color=(0.2, 0.4, 0.6))
    plt.plot(kList, recallList, label="Recall", color=(0.2, 0.7, 0.5))
    plt.plot(kList, f1List, label="F1", color=(0.7, 0.9, 0.1))
    plt.legend(loc=4)
    plt.ylim(0, 1)
    plt.title('Classification quality char for all k')
    plt.xlabel('k')

def plotClassificationQualityForMultiLabels(listOfTrueLabels, listOfPredictedLabels, listOfK):
    accList = []
    kList = []
    for trueLabels, predictedLabels, k in zip(listOfTrueLabels, listOfPredictedLabels, listOfK):
        accList.append(accuracy_score(trueLabels, predictedLabels))
        kList.append(k)

    order = np.argsort(kList)
    accList = np.array(accList)[order]
    kList = np.array(kList)[order]

    plt.plot(kList, accList, label="Precision", color=(0.2, 0.4, 0.6))
    plt.legend(loc=4)
    plt.ylim(0, 1)
    plt.title('Accuracy char for all k')
    plt.xlabel('k')


def extraktNumberOfK(arg):
    head, tail = os.path.split(arg)
    tail = tail.replace('k_', '').replace('.txt', '')
    return int(tail)


def extraktLastDirFromPath(arg):
    head, tail = os.path.split(arg)
    return tail


def isCorrectFileName(file):
    if '.txt' not in file:
        return False
    elif 'config' in file:
        return False
    return True


def createDocForOneMetric(dir, label_mode):
    metrics = extraktLastDirFromPath(dir)
    document = Document()
    header = document.add_heading('Report: ', 1)

    listOfTrueLabels = []
    listOfPredictedLabels = []
    listOfK = []

    for r, d, f in os.walk(dir):
        header.add_run(metrics + ' mertics').bold = True
        for file in f:
            if isCorrectFileName(file):
                pathToFileWithClassificationData = os.path.abspath(dir)

                k = extraktNumberOfK(file)
                trueLabels, predictedLabels = extractClassificationCases(pathToFileWithClassificationData + "\\" + file)
                plot_confusion_matrix(trueLabels, predictedLabels, classes=unique_labels(trueLabels, predictedLabels),
                                      normalize=True,
                                      title='Normalized confusion matrix for ' + str(k) + ' k')
                memfile = StringIO()
                plt.savefig(memfile)
                document.add_paragraph('Confusion matrix for ' + str(k) + ' k.', style='List Number')
                document.add_picture(memfile, width=Inches(5))
                memfile.close()
                plt.close()

                listOfTrueLabels.append(trueLabels)
                listOfPredictedLabels.append(predictedLabels)
                listOfK.append(k)

    if label_mode == "single":
        plotClassificationQualityForSingleLabel(listOfTrueLabels, listOfPredictedLabels, listOfK)
    elif label_mode == "multi":
        plotClassificationQualityForMultiLabels(listOfTrueLabels, listOfPredictedLabels, listOfK)
    memfile = StringIO()
    plt.savefig(memfile)
    document.add_paragraph('Classification quality char for all k.', style='List Number')
    document.add_picture(memfile, width=Inches(6))
    memfile.close()
    document.save('report_' + metrics + '.docx')


def main():
    try:
        opts, args = getopt.getopt(sys.argv[1:], "f:d:r:sm", ["filePath=", "dir=", "report=","singleLabel=","multiLabel="])
    except getopt.GetoptError:
        sys.exit(2)
    for opt, arg in opts:
        if opt in ("-f", "--file"):
            filePath = arg
            pathToFileWithClassificationData = os.path.abspath(filePath)
            trueLabels, predictedLabels = extractClassificationCases(pathToFileWithClassificationData)

            plot_confusion_matrix(trueLabels, predictedLabels, classes=unique_labels(trueLabels, predictedLabels),
                                  normalize=True,
                                  title='Normalized confusion matrix')
            plt.show()


        elif opt in ("-d", "--dir"):
            dir = arg

            for r, d, f in os.walk(dir):
                for file in f:
                    if isCorrectFileName(file):
                        pathToFileWithClassificationData = os.path.abspath(dir)
                        trueLabels, predictedLabels = extractClassificationCases(
                            pathToFileWithClassificationData + "\\" + file)

                        print(file)

        elif opt in ("-r", "--report"):
            try:
                if ('-s','') in opts:
                    createDocForOneMetric(arg, "single")
                elif ('-m','') in opts:
                    createDocForOneMetric(arg, "multi")
                else:
                    raise ValueError
            except ValueError:
                print 'must be defined additional option -s when single label or -m when multilabel', arg


if __name__ == "__main__":
    main()
